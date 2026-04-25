import io
import re
import os
import unicodedata
import os
from fastapi import FastAPI, UploadFile, File, HTTPException, Query, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import PyPDF2
import docx
import spacy
from github_client import GitHubExtractor
from tcfe_engine import calculate_tcfe
from langchain_engine import embed_document, evaluate_candidate
from graph_engine import generate_knowledge_graph, calculate_gap_analysis
import asyncio
from auth import auth_router, get_current_user

app = FastAPI(title="TalentGraph AI - Resume Parser API")

app.include_router(auth_router)

from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse

# --- 🚨 CORS MIDDLEWARE FIX (For Frontend Integration) ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Wildcard works ONLY if allow_credentials=False
    allow_credentials=False, # We use Bearer tokens, not cookies, so this can be False
    allow_methods=["*"], 
    allow_headers=["*"], 
)

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc):
    print(f"OMG 422 ERROR: {exc}")
    print(f"BODY WAS: {await request.body()}")
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors(), "body": str(exc)},
    )

# --- HEALTH CHECK (For React Status Bar) ---
@app.get("/health")
async def health_check():
    """Frontend polling endpoint for the Status Bar"""
    return {"status": "online", "vector_db": "connected", "tcfe": "active"}

# Load spaCy model for PII Redaction
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    raise RuntimeError(
        "spaCy model 'en_core_web_sm' not found. "
        "Please install it using: python -m spacy download en_core_web_sm"
    )

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extracts text from a PDF file using PyPDF2 and scrubs invisible characters."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        # Allow all unicode, just strip control characters except newline
        return "".join(ch for ch in text if unicodedata.category(ch)[0] != "C" or ch == '\n')
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extracts text from a Word document using python-docx."""
    try:
        import unicodedata
        doc = docx.Document(io.BytesIO(file_content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return "".join(ch for ch in text if unicodedata.category(ch)[0] != "C" or ch == '\n')
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

def extract_github_from_docx_hyperlinks(file_content: bytes) -> Optional[str]:
    """
    Strategy 3: Scans embedded hyperlink relationships inside a DOCX file.
    These are stored in document.xml.rels, NOT as visible text, so regex on
    extracted text will never find them.
    """
    try:
        doc = docx.Document(io.BytesIO(file_content))
        # Iterate all relationships in the document part
        for rel in doc.part.rels.values():
            target = rel.target_ref or ""
            if "github.com" in target.lower():
                # Extract username from the URL
                match = re.search(r'github\.com/([a-zA-Z0-9_-]+)', target, re.IGNORECASE)
                if match:
                    username = match.group(1)
                    # Filter out known non-profile paths
                    if username.lower() not in {"orgs", "explore", "topics", "trending", "login"}:
                        return f"https://github.com/{username}"
    except Exception:
        pass
    return None

def extract_github_from_pdf_annotations(file_content: bytes) -> Optional[str]:
    """
    Strategy 4: Scans PDF page annotations (clickable hyperlinks) for a GitHub URL.
    PyPDF2 can read /URI entries from annotation dictionaries even when the
    link text is just a word like 'GitHub' with no visible URL.
    """
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            if "/Annots" not in page:
                continue
            annotations = page["/Annots"]
            if annotations is None:
                continue
            for annot in annotations:
                obj = annot.get_object() if hasattr(annot, 'get_object') else annot
                if obj.get("/Subtype") == "/Link":
                    action = obj.get("/A")
                    if action and action.get("/URI"):
                        uri = str(action["/URI"])
                        if "github.com" in uri.lower():
                            match = re.search(r'github\.com/([a-zA-Z0-9_-]+)', uri, re.IGNORECASE)
                            if match:
                                username = match.group(1)
                                if username.lower() not in {"orgs", "explore", "topics", "trending", "login"}:
                                    return f"https://github.com/{username}"
    except Exception:
        pass
    return None


def extract_github_url(text: str, file_content: bytes = None, file_type: str = None) -> Optional[str]:
    """
    4-strategy GitHub profile extraction chain.
    Strategy 1: Full URL in plain text  (e.g. https://github.com/torvalds)
    Strategy 2: Labeled username in plain text (e.g. 'GitHub: torvalds' or '@torvalds')
    Strategy 3: Embedded DOCX hyperlink in document relationships
    Strategy 4: PDF annotation /URI hyperlink object
    Returns a normalized https://github.com/<username> URL or None.
    """
    # --- Strategy 1: Full or partial URL in extracted text ---
    match = re.search(r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_-]+)', text, re.IGNORECASE)
    if match:
        username = match.group(1)
        if username.lower() not in {"orgs", "explore", "topics", "trending", "login"}:
            return f"https://github.com/{username}"

    # --- Strategy 2: Labeled username patterns ---
    # Forward:  "GitHub: torvalds", "GitHub - torvalds", "GitHub | torvalds", "GitHub @torvalds"
    labeled = re.search(
        r'github\s*[:\-|@\s]+\s*@?([a-zA-Z0-9_-]{1,39})(?:\s|$|,|\|)',
        text,
        re.IGNORECASE
    )
    if labeled:
        username = labeled.group(1).strip()
        if username.lower() not in {"com", "profile", "orgs", "explore"}:
            return f"https://github.com/{username}"

    # Reverse: "@torvalds (GitHub)", "@torvalds - GitHub", "torvalds | GitHub"
    reverse = re.search(
        r'@?([a-zA-Z0-9_-]{1,39})\s*[\(\-\|,]?\s*github',
        text,
        re.IGNORECASE
    )
    if reverse:
        username = reverse.group(1).strip()
        if username.lower() not in {"com", "profile", "orgs", "explore", "on", "via", "at"}:
            return f"https://github.com/{username}"


    # --- Strategy 3: DOCX embedded hyperlinks ---
    if file_content and file_type == "docx":
        result = extract_github_from_docx_hyperlinks(file_content)
        if result:
            return result

    # --- Strategy 4: PDF annotation hyperlinks ---
    if file_content and file_type == "pdf":
        result = extract_github_from_pdf_annotations(file_content)
        if result:
            return result

    return None


def redact_entities(text: str) -> str:
    """Uses spaCy to detect and redact Person Names and Organizations for blind screening."""
    doc = nlp(text)
    redacted_text = text
    
    # Sort entities in reverse order to avoid index shifting when replacing text
    entities = sorted(
        [ent for ent in doc.ents if ent.label_ in ["PERSON", "ORG", "GPE"]],
        key=lambda e: e.start_char,
        reverse=True
    )
    
    for ent in entities:
        start = ent.start_char
        end = ent.end_char
        replacement = f"[{ent.label_}]"
        redacted_text = redacted_text[:start] + replacement + redacted_text[end:]
        
    return redacted_text

async def llm_sterilize_resume(raw_text: str) -> str:
    """
    Uses Llama 3.2 to rewrite the resume and eliminate proxy bias.
    Falls back to spaCy-based redact_entities on failure or timeout.
    """
    prompt = f"""You are a Bias Sterilization Engine for a strict meritocratic hiring platform. 
Your objective is to rewrite the provided resume achievements to eliminate all explicit and implicit "Proxy Bias."

RULES:
1. REMOVE all names of individuals, companies, universities, fraternities/sororities, non-profits, and geographic locations.
2. REPLACE them with generic functional equivalents (e.g., change "Stanford University" to "[Tier 1 Academic Institution]", change "Google" to "[Enterprise Tech Company]", change "Women in Tech" to "[Diversity Organization]").
3. REMOVE all dates, graduation years, and lengths of tenure to prevent ageism.
4. PRESERVE 100% of the technical context, frameworks, code deployment metrics, and business impact. 
5. OUTPUT ONLY THE STERILIZED RESUME TEXT. DO NOT OUTPUT ANY CONVERSATIONAL TEXT, NOTES, OR INTRODUCTIONS.

Input Text: {raw_text}
Sterilized Output:"""

    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "llama3.2",
        "prompt": prompt,
        "stream": False
    }
    
    try:
        import httpx
        async with httpx.AsyncClient(timeout=60.0) as client:
            res = await client.post(url, json=payload)
            if res.status_code == 200:
                data = res.json()
                response_text = data.get("response", "").strip()
                if response_text:
                    return response_text
    except Exception as e:
        print(f"LLM Sterilization Error: {e}. Falling back to spaCy.")
    
    # Fallback to legacy spaCy engine
    return redact_entities(raw_text)


@app.post("/parse-resume")
async def parse_resume(file: UploadFile = File(...), current_user: str = Depends(get_current_user)):
    """
    Phase 1 & 2: Parses PDF or DOCX, redacts PII, checks for technical relevance, 
    and extracts live GitHub metrics (TCFE).
    """
    filename_lower = file.filename.lower()
    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF and Word (.docx) documents are accepted.")
    
    try:
        content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail="Could not read the uploaded file.")
    
    # 1. Extract raw text based on file extension
    try:
        if filename_lower.endswith(".pdf"):
            raw_text = extract_text_from_pdf(content)
        else:
            raw_text = extract_text_from_docx(content)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract any text from the document.")
    
    # 2. Redact entities FIRST (Blind Screening) via LLM Bias Sterilization
    sanitized_text = await llm_sterilize_resume(raw_text)
    
    # 3. Relevance Gating check
    tech_keywords = ["engineer", "developer", "software", "data", "code", "tech"]
    text_lower = sanitized_text.lower()
    is_technical = any(kw in text_lower for kw in tech_keywords)
    
    if not is_technical:
        return {
            "github_url": None,
            "tcfe_metrics": None,
            "sanitized_text": sanitized_text,
            "message": "Non-technical role detected. Skipping TCFE."
        }
    
    # 4. Proceed with GitHub extraction for technical roles
    # Extract from raw_text to ensure the LLM didn't accidentally remove the URL.
    # Also pass raw bytes + file type to scan embedded hyperlinks (DOCX rels / PDF annotations).
    file_type = "pdf" if filename_lower.endswith(".pdf") else "docx"
    github_url = extract_github_url(raw_text, file_content=content, file_type=file_type)

    
    tcfe_metrics = None
    if github_url:
        try:
            # Safely get the username from the raw URL
            username = github_url.rstrip('/').split('/')[-1]
            if username:
                print(f"Fetching GitHub API for user: {username}")
                github_token = os.getenv("GITHUB_TOKEN")
                extractor = GitHubExtractor(token=github_token)
                top_repos = await extractor.get_top_repositories(username, limit=1)
                
                if top_repos:
                    repo_info = top_repos[0]
                    repo_name = repo_info.get("name")
                    repo_owner = repo_info.get("owner", {}).get("login")
                    repo_created_at = repo_info.get("created_at")
                    
                    if repo_name and repo_owner and repo_created_at:
                        commits = await extractor.get_recent_commits(repo_owner, repo_name)
                        # Run the blocking LLM logic in a separate thread
                        tcfe_metrics = await asyncio.to_thread(calculate_tcfe, commits, repo_created_at)
        except Exception:
            tcfe_metrics = None
    
    return {
        "github_url": github_url,
        "tcfe_metrics": tcfe_metrics,
        "sanitized_text": sanitized_text,
        "message": "Technical role processed successfully."
    }

# --- Models and Endpoints for Phase 3 (LLM) & Phase 4 (Graph) ---

class EmbedStoreRequest(BaseModel):
    candidate_id: str
    sanitized_text: str
    metadata: Optional[Dict[str, Any]] = None

class EvaluateRequest(BaseModel):
    candidate_id: str
    job_description: str = ""
    pow_data: Optional[Dict[str, Any]] = None
    role_id: Optional[int] = None

@app.post("/embed-store")
async def embed_store(request: EmbedStoreRequest, current_user: str = Depends(get_current_user)):
    """Chunks and embeds the sanitized resume into ChromaDB."""
    try:
        await asyncio.to_thread(embed_document, request.candidate_id, request.sanitized_text, request.metadata or {})
        return {"status": "success", "message": f"Document embedded successfully for candidate {request.candidate_id}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error embedding document: {str(e)}")

@app.post("/evaluate-candidate")
async def evaluate_candidate_endpoint(request: EvaluateRequest, current_user: str = Depends(get_current_user)):
    """
    The Mega-Endpoint: Queries ChromaDB, runs Llama 3.2, applies TCFE momentum math, 
    and generates the D3 topology graph in one payload.
    """
    try:
        active_pow_data = request.pow_data
        if not active_pow_data:
            active_pow_data = {}

        # Fetch Role Context from DB if provided
        role_context = ""
        if request.role_id:
            import sqlite3
            from auth import DB_FILE
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT title, description FROM roles WHERE id = ?", (request.role_id,))
            row = cursor.fetchone()
            conn.close()
            if row:
                role_context = f"Role Title: {row[0]}\nRole Context: {row[1]}"

        # Calculate Advanced GitHub PoW math
        pow_results = None
        if active_pow_data.get("github_user"):
            from pow_scoring import calculate_pow_score
            pow_results = await calculate_pow_score(active_pow_data.get("github_user"), request.job_description)

        # 1. Run blocking LLM evaluation
        llm_evaluation = await asyncio.to_thread(
            evaluate_candidate, 
            request.candidate_id, 
            request.job_description, 
            active_pow_data,
            role_context,
            pow_results
        )

        print("\n=== RAW LLM OUTPUT ===")
        print(llm_evaluation)
        print("======================\n")
        
        # 2. Extract skills
        candidate_skills = llm_evaluation.get("extracted_candidate_skills", [])
        jd_skills = llm_evaluation.get("extracted_jd_skills", [])
        
        # 3. Generate Graph and Gap Analysis
        graph_data = generate_knowledge_graph(candidate_skills, jd_skills)
        gap_analysis = calculate_gap_analysis(candidate_skills, jd_skills)
        
        # New PoW Pipeline directly uses the total score which includes the deterministic PoW override
        final_weighted_score = llm_evaluation.get("total_score", 0.0)
        
        # Format the temporal velocity block for frontend display compatibility
        burst_triggered = pow_results.get("burst_triggered", False) if pow_results else False
        temporal_status = "Accelerated" if burst_triggered else "Stable"

        # 4. Build the unified JSON payload
        unified_response = {
            "candidate_id": request.candidate_id,
            "scores": {
                "semantic_skill_score_40": llm_evaluation.get("semantic_skill_score_40"),
                "pow_depth_score_30": llm_evaluation.get("pow_depth_score_30"),
                "experience_score_15": llm_evaluation.get("experience_score_15"),
                "keyword_score_15": llm_evaluation.get("keyword_score_15"),
                "total_score": final_weighted_score
            },
            "temporal_velocity": {
                "status": temporal_status,
                "multiplier_applied": pow_results.get("velocity_component", 0) if pow_results else 0.0,
                "final_weighted_score": round(final_weighted_score, 2),
                "burst_flag": burst_triggered,
                "raw_pow_score_100": pow_results.get("pow_score", 0) if pow_results else 0.0
            },
            "explanation": llm_evaluation.get("xai_explanation"),
            "gap_analysis": gap_analysis,
            "graph_data": graph_data
        }
        
        return unified_response
    except ValueError as ve:
        raise HTTPException(status_code=404, detail=str(ve))
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print("ERROR IN EVALUATE CANDIDATE:\n", error_details)
        raise HTTPException(status_code=500, detail=f"Error scoring candidate: {str(e)}\n\n{error_details}")


class SkillsPayload(BaseModel):
    candidate_skills: List[str]
    jd_skills: List[str]

@app.post("/graph-data")
async def get_graph_data(payload: SkillsPayload):
    """Returns a D3-compatible JSON object generated by NetworkX."""
    try:
        graph_data = generate_knowledge_graph(payload.candidate_skills, payload.jd_skills)
        return graph_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating graph data: {str(e)}")

@app.post("/gap-analysis")
async def gap_analysis_endpoint(payload: SkillsPayload):
    """Analyzes gaps between candidate skills and the JD using set mathematics."""
    try:
        analysis_data = calculate_gap_analysis(payload.candidate_skills, payload.jd_skills)
        return analysis_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error calculating gap analysis: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)